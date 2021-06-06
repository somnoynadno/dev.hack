from docx import Document
from docx.shared import Inches
import os.path
import random
import string

default_upload_folder = '/tmp'
UPLOAD_FOLDER = os.environ.get("UPLOAD_FOLDER", default_upload_folder)

filename_length = 16

def create_operation_status_docx(account_from, account_to, currency_code_from, currency_code_to, currency_value_from, currency_value_to, status, comission, filename):
    document = Document()
    t = document.add_heading('ПСБ-Банк', 0)
    p = t.add_run()
    p.add_picture('logo.png', width=Inches(0.25))
    document.add_paragraph(f'Счет списания: {account_from}')
    document.add_paragraph(f'Счет получения: {account_to}')
    document.add_paragraph(f'Списано: {currency_value_from} {currency_code_from}')
    document.add_paragraph(f'Зачислено: {currency_value_to} {currency_code_to}')
    document.add_paragraph(f'Комиссия: {comission}')
    document.add_paragraph(f'Статус операции: {status}')
    filename = os.path.join(UPLOAD_FOLDER, filename)
    document.save(filename)

def generate_random_filename():
    return ''.join(random.choices(string.ascii_letters + string.digits, k=filename_length))+'.docx'

