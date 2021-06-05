from docx import Document
from docx.shared import Inches
import os.path

default_upload_folder = '/tmp'
UPLOAD_FOLDER = os.environ.get("UPLOAD_FOLDER", default_upload_folder)

def create_operation_status_docx(account_from, account_to, currency_code, currency_value, status):
    document = Document()
    document.add_heading('ПСБ-Банк', 0)
    document.add_paragraph(f'Счет списания: {account_from}')
    document.add_paragraph(f'Счет получения: {account_to}')
    document.add_paragraph(f'Сумма: {currency_value} {currency_code}')
    document.add_paragraph(f'Статус операции: {status}')

    filename = os.path.join(UPLOAD_FOLDER, "test.docx")
    document.save(filename)

create_operation_status_docx('2112', '3232', 'EUR', 12, "Отменена")
    

